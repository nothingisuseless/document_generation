import os
import streamlit as st
from langchain.document_loaders import DirectoryLoader, PyPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA, LLMChain
from langchain.prompts import PromptTemplate
from langchain.llms import Ollama
from docx import Document
from docx.shared import Pt

# Set page config
st.set_page_config(page_title="Document Generator (FRS/Design)", layout="centered")
st.title("🛠️ Generate Project Documents from FRS/Design Examples")

@st.cache_resource(show_spinner="Indexing FRS/Design Docs...")
def load_vectorstore(pdf_dir="pdfs/"):
    loader = DirectoryLoader(pdf_dir, glob="**/*.pdf", loader_cls=PyPDFLoader)
    documents = loader.load()
    splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_documents(documents)

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(chunks, embeddings)

    return vectorstore

# Load VectorStore and Model
vectorstore = load_vectorstore()
llm = Ollama(model="llama3")

# User input for document generation
doc_type = st.selectbox("Select document type to generate:", ["Functional Requirement Specification (FRS)", "Design Document"])
project_description = st.text_area("Describe your project requirements:")

from docx.enum.style import WD_STYLE_TYPE

def save_to_word(title, content, filename="Generated_Document.docx"):
    doc = Document()
    doc.add_heading(title, 0)

    for section in content.split("\n\n"):
        if section.strip():
            lines = section.strip().split("\n", 1)
            if len(lines) == 2:
                heading, body = lines
                doc.add_heading(heading.strip(), level=1)
                doc.add_paragraph(body.strip())
            else:
                doc.add_paragraph(section.strip())

    doc.save(filename)


# Prompt Template
prompt_template = """
You are a senior software architect.

Generate a complete and professionally structured {doc_type} document based on the following project description:

"{project_description}"

Refer to the following relevant context from similar documents:
"{context}"

Structure the output with clearly defined sections.

If the document type is:
- "Functional Requirement Specification (FRS)", use this structure:
  1. Introduction
  2. Purpose
  3. Scope
  4. Definitions, Acronyms, Abbreviations
  5. Functional Requirements
  6. Non-Functional Requirements
  7. Assumptions and Constraints

If the document type is:
- "Design Document", use this structure:
  1. Overview
  2. Architecture Description
  3. Module Design
  4. Database Design
  5. Interface Design
  6. Security Considerations
  7. Assumptions and Limitations

Write the document in a formal, technical tone with headings for each section.
"""


if st.button("Generate Document"):
    if not project_description.strip():
        st.error("Please provide a project description.")
    else:
        with st.spinner("Generating..."):

            # Use Retriever to get relevant chunks from past docs
            retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
            relevant_docs = retriever.get_relevant_documents(project_description)
            context = "\n\n".join([doc.page_content for doc in relevant_docs])

            # Format prompt
            prompt = PromptTemplate(
                input_variables=["project_description", "context", "doc_type"],
                template=prompt_template
            )
            chain = LLMChain(llm=llm, prompt=prompt)
            result = chain.run({
                "project_description": project_description,
                "context": context,
                "doc_type": doc_type
            })

            st.success(f"{doc_type} Generated!")
            st.write(result)

            # Save and allow download
            save_to_word(doc_type, result)
            st.download_button("📄 Download Generated Document", data=open("Generated_Document.docx", "rb"), file_name="Generated_Document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
